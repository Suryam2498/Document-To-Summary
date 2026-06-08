"""
doc_app.py — Standalone Document Q&A Assistant
Run: uvicorn doc_app:app --port 8001 --reload
"""

from fastapi import FastAPI, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import PlainTextResponse, StreamingResponse
from pydantic import BaseModel
import pdfplumber
import io, uuid, json
import os
from pathlib import Path
from datetime import datetime
from groq import Groq
from docx import Document as DocxDoc
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    from docx import Document as DocxDocument
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

# ── Config ────────────────────────────────────────────────────────────────────
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "").strip()
LLM_MODEL    = "llama-3.3-70b-versatile"
MAX_DOCS     = 10
MAX_CHARS    = 10000   # chars per doc sent to LLM

groq_client  = Groq(api_key=GROQ_API_KEY) if GROQ_API_KEY else None

# ── App ───────────────────────────────────────────────────────────────────────
app = FastAPI(title="DocChat Assistant")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

# ── In-memory stores ──────────────────────────────────────────────────────────
doc_store: dict = {}   # {doc_id: {name, text, summary, pages, type, chars, uploaded_at}}
chat_history: list = []  # [{role, content, timestamp}]

# ── Text extraction ───────────────────────────────────────────────────────────
def extract_pdf(content: bytes) -> tuple[str, int]:
    text, pages = "", 0
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        pages = len(pdf.pages)
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"
    return text, pages

def extract_docx(content: bytes) -> tuple[str, int]:
    doc = DocxDocument(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    pages = max(1, len(paragraphs) // 25)
    return text, pages

# ── LLM helpers ───────────────────────────────────────────────────────────────
def summarize(name: str, text: str) -> str:
    if groq_client is None:
        raise RuntimeError("Missing GROQ_API_KEY environment variable.")

    resp = groq_client.chat.completions.create(
        model=LLM_MODEL,
        messages=[{
            "role": "user",
            "content": (
                f"Summarize the following document in 4-5 concise bullet points. "
                f"Be specific about key topics, data, and findings.\n\n"
                f"Document: {name}\n\n{text[:6000]}"
            )
        }],
        temperature=0.3,
        max_tokens=500,
    )
    return resp.choices[0].message.content

def build_context() -> str:
    parts = []
    for doc in doc_store.values():
        parts.append(f"=== {doc['name']} ===\n{doc['text'][:MAX_CHARS]}")
    return "\n\n".join(parts)

# ── Routes ────────────────────────────────────────────────────────────────────
class AskRequest(BaseModel):
    question: str

@app.post("/upload")
async def upload(file: UploadFile = File(...)):
    if len(doc_store) >= MAX_DOCS:
        return {"error": f"Maximum {MAX_DOCS} documents reached. Remove one first."}

    fname = file.filename.lower()
    if fname.endswith(".pdf"):
        ftype = "PDF"
    elif fname.endswith(".docx"):
        if not DOCX_SUPPORT:
            return {"error": "python-docx not installed. Run: pip install python-docx"}
        ftype = "DOCX"
    else:
        return {"error": "Only PDF and DOCX files are supported."}

    content = await file.read()

    try:
        text, pages = extract_pdf(content) if ftype == "PDF" else extract_docx(content)
        if not text.strip():
            return {"error": "Could not extract text from this document."}

        doc_id  = str(uuid.uuid4())[:8]
        summary = summarize(file.filename, text)

        doc_store[doc_id] = {
            "name":         file.filename,
            "text":         text,
            "summary":      summary,
            "pages":        pages,
            "type":         ftype,
            "chars":        len(text),
            "uploaded_at":  datetime.now().strftime("%H:%M:%S"),
        }

        return {
            "doc_id":  doc_id,
            "name":    file.filename,
            "type":    ftype,
            "pages":   pages,
            "chars":   len(text),
            "summary": summary,
        }
    except Exception as e:
        return {"error": str(e)}


@app.post("/ask")
async def ask(req: AskRequest):
    if not doc_store:
        return {"error": "No documents uploaded yet."}

    if groq_client is None:
        return {"error": "Missing GROQ_API_KEY environment variable."}

    context   = build_context()
    doc_names = [d["name"] for d in doc_store.values()]

    history_msgs = [{"role": h["role"], "content": h["content"]} for h in chat_history[-12:]]

    system = (
        "You are a document analysis assistant. Answer questions based ONLY on the provided documents.\n"
        "If the answer is not in the documents, say so clearly.\n"
        "Be concise and accurate. Cite which document you're referencing when relevant.\n\n"
        f"Documents: {', '.join(doc_names)}\n\n"
        f"Content:\n{context}"
    )

    messages = [{"role": "system", "content": system}] + history_msgs + [{"role": "user", "content": req.question}]

    resp = groq_client.chat.completions.create(
        model=LLM_MODEL,
        messages=messages,
        temperature=0.2,
        max_tokens=1000,
    )
    answer = resp.choices[0].message.content
    ts     = datetime.now().strftime("%H:%M:%S")

    chat_history.append({"role": "user",      "content": req.question, "timestamp": ts})
    chat_history.append({"role": "assistant", "content": answer,       "timestamp": ts})

    return {"answer": answer, "sources": doc_names}


@app.get("/docs")
def list_docs():
    return {"docs": [{
        "id":          k,
        "name":        v["name"],
        "type":        v["type"],
        "pages":       v["pages"],
        "chars":       v["chars"],
        "summary":     v["summary"],
        "uploaded_at": v["uploaded_at"],
    } for k, v in doc_store.items()]}


@app.delete("/docs/{doc_id}")
def remove_doc(doc_id: str):
    if doc_id in doc_store:
        name = doc_store.pop(doc_id)["name"]
        return {"removed": name}
    return {"error": "Document not found"}


@app.delete("/clear")
def clear_all():
    doc_store.clear()
    chat_history.clear()
    return {"status": "cleared"}


@app.get("/history")
def get_history():
    return {"history": chat_history}


@app.get("/download")
def download():
    doc = DocxDoc()

    # Title
    title = doc.add_heading("DocChat — Conversation History", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)

    # Date line
    date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    date_para.runs[0].font.size = Pt(9)
    date_para.runs[0].font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)
    doc.add_paragraph("")

    if not chat_history:
        doc.add_paragraph("No conversation history yet.")
    else:
        for h in chat_history:
            if h["role"] == "user":
                p = doc.add_paragraph()
                p.add_run(f"[{h['timestamp']}] You:  ").bold = True
                p.add_run(h["content"])
            else:
                p = doc.add_paragraph()
                label = p.add_run(f"[{h['timestamp']}] Assistant:  ")
                label.bold = True
                label.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)
                p.add_run(h["content"])
            doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"docchat-{datetime.now().strftime('%Y-%m-%d')}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@app.get("/health")
def health():
    return {
        "status":       "ok",
        "docs_loaded":  len(doc_store),
        "docx_support": DOCX_SUPPORT,
    }


# ── Serve static ──────────────────────────────────────────────────────────────
app.mount("/", StaticFiles(directory=str(Path(__file__).parent), html=True), name="static")
